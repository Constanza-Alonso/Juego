from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "INFORME_PROYECTO.md"
OUTPUT = ROOT / "Informe_Shadow_Beat_Fragmentos_de_Luz.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 99, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"

SPANISH_REPLACEMENTS = {
    "Titulo": "Título",
    "titulo": "título",
    "ritmico": "rítmico",
    "precision": "precisión",
    "logica": "lógica",
    "automatico": "automático",
    "automaticamente": "automáticamente",
    "mecanica": "mecánica",
    "mecanicas": "mecánicas",
    "progresion": "progresión",
    "narrativa": "narrativa",
    "puntuacion": "puntuación",
    "accion": "acción",
    "direccion": "dirección",
    "fisicas": "físicas",
    "fisica": "física",
    "generacion": "generación",
    "genero": "género",
    "Tambien": "También",
    "tambien": "también",
    "Nucleo": "Núcleo",
    "nucleo": "núcleo",
    "destruyo": "destruyó",
    "disperso": "dispersó",
    "quedo": "quedó",
    "ultimo": "último",
    "Caracteristicas": "Características",
    "caracteristicas": "características",
    "obstaculos": "obstáculos",
    "obstaculo": "obstáculo",
    "reintento": "reintento",
    "version": "versión",
    "basicos": "básicos",
    "basicas": "básicas",
    "basica": "básica",
    "camara": "cámara",
    "camaras": "cámaras",
    "menu": "menú",
    "menus": "menús",
    "numero": "número",
    "bonificacion": "bonificación",
    "finalizacion": "finalización",
    "funcion": "función",
    "Funcion": "Función",
    "opcion": "opción",
    "opciones": "opciones",
    "seleccion": "selección",
    "dificultad": "dificultad",
    "mas": "más",
    "todavia": "todavía",
    "moviles": "móviles",
    "movil": "móvil",
    "caidas": "caídas",
    "caida": "caída",
    "vacio": "vacío",
    "segun": "según",
    "pequeno": "pequeño",
    "traves": "través",
    "atmosfera": "atmósfera",
    "triangulares": "triangulares",
    "luminosos": "luminosos",
    "recolectables": "recolectables",
    "codigo": "código",
    "credito": "crédito",
    "creditos": "créditos",
    "musica": "música",
    "particulas": "partículas",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_keep_with_next(paragraph, enabled=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if enabled and keep is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not enabled and keep is not None:
        p_pr.remove(keep)


def add_formatted_runs(paragraph, text: str, bold_default=False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = polish_text(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
        else:
            run.text = polish_text(part)
            run.bold = bold_default


def polish_text(text: str) -> str:
    for plain, accented in SPANISH_REPLACEMENTS.items():
        text = re.sub(rf"\b{re.escape(plain)}\b", accented, text)
    return text


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Shadow Beat: Fragmentos de Luz"
    header.style = doc.styles["Header"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Informe de proyecto")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Shadow Beat: Fragmentos de Luz")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    sub = subtitle.add_run("Informe del proyecto de videojuego")
    sub.font.size = Pt(14)
    sub.font.color.rgb = BLUE

    facts = [
        ("Motor", "Unity 2022.3 LTS"),
        ("Genero", "Plataformas ritmico / arcade de precision"),
        ("Proyecto", "Juego 2D con avance automatico, niveles, portales y cristales"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    for label, value in facts:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].runs[0].bold = True

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(22)
    note.paragraph_format.space_after = Pt(8)
    add_formatted_runs(
        note,
        "Este documento resume el concepto, las mecanicas, el estado de desarrollo y los elementos implementados en la primera etapa funcional del juego.",
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_formatted_runs(p, text)


def add_table_from_markdown(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_formatted_runs(p, value)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for run in p.runs:
                    run.bold = True
        if r_idx == 0:
            set_repeat_table_header(table.rows[0])

    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    add_formatted_runs(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    add_formatted_runs(p, text)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)

    lines = markdown.splitlines()
    pending_table: list[str] = []
    skip_title_section = False

    add_callout(
        doc,
        "Estado actual: version funcional inicial con menu de seleccion, siete niveles, mecanicas principales, portales, cristales, sistema de reintento y finalizacion de niveles.",
    )

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if pending_table and (not stripped.startswith("|")):
            add_table_from_markdown(doc, pending_table)
            pending_table = []

        if not stripped:
            continue

        if stripped.startswith("|"):
            pending_table.append(stripped)
            continue

        if stripped == "## Titulo":
            skip_title_section = True
            continue
        if skip_title_section:
            if stripped.startswith("## "):
                skip_title_section = False
            else:
                continue

        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(polish_text(stripped[3:]), level=1)
            set_paragraph_keep_with_next(p)
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(polish_text(stripped[4:]), level=2)
            set_paragraph_keep_with_next(p)
            continue
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            continue
        if re.match(r"^\d+\.\s+", stripped):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)

    if pending_table:
        add_table_from_markdown(doc, pending_table)

    doc.core_properties.title = "Informe Shadow Beat: Fragmentos de Luz"
    doc.core_properties.subject = "Proyecto de videojuego en Unity"
    doc.core_properties.author = "Constanza Alonso Winter"
    doc.core_properties.comments = "Informe generado para documentar la primera etapa funcional del proyecto."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
